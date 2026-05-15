import os
import re
from docx import Document as DocxDocument
from langchain_core.documents import Document

def parse_legal_document(file_path):
    doc = DocxDocument(file_path)
    chunks = []

    document_name = "Chưa xác định"
    # Thử lấy tên văn bản từ vài dòng đầu tiên
    for i in range(min(5, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if text and not re.search(r'^(Chương|Mục|Điều|CỘNG HÒA|Độc lập)', text, re.I):
            document_name = text
            break

    current_chapter = "0"
    current_section = "0"
    current_article = "0"
    current_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue

        chapter_match = re.search(r'^(Chương\s+([IVXLCDM\d]+))', text, re.I)
        section_match = re.search(r'^(Mục\s+(\d+))', text, re.I)
        article_match = re.search(r'^(Điều\s+(\d+\w*))', text, re.I)

        if article_match:
            if current_article != "0" and current_content:
                chunk_metadata = {
                    "source": os.path.basename(file_path),
                    "document_name": document_name,
                    "chapter": current_chapter,
                    "section": current_section,
                    "article": current_article,
                }
                chunks.append(Document(page_content="\n".join(current_content), metadata=chunk_metadata))

            current_article = article_match.group(2)
            current_content = [text]

        elif chapter_match:
            current_chapter = chapter_match.group(2); current_content.append(text)
        elif section_match:
            current_section = section_match.group(2); current_content.append(text)
        elif current_article != "0":
            current_content.append(text)

    if current_article != "0" and current_content:
        chunk_metadata = {
            "source": os.path.basename(file_path),
            "document_name": document_name,
            "chapter": current_chapter,
            "section": current_section,
            "article": current_article,
        }
        chunks.append(Document(page_content="\n".join(current_content), metadata=chunk_metadata))

    return chunks